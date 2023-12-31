from flask import Flask, render_template, request, session, send_file, send_from_directory
from flask_cors import CORS, cross_origin
import ai21
from docx import Document
from docx.shared import Inches
import requests
import redis
import os

app = Flask(__name__)
cors = CORS(app)
app.config['CORS_HEADERS'] = 'Content-Type'

def get_transcription(url: str) -> str:
    value = r.get(url)
    if value is None:
        return ""
    else:
        return value.decode("utf-8")
    
def append_transcription(url: str, transcript: str):
    # if url in r: append to existing
    # else: create new
    existing = r.get(url)
    if existing is None:
        r.set(url, transcript)
    else:
        existing = existing.decode("utf-8")
        r.set(url, existing + ' ' + transcript)

def ask(transcript, question):
    headers = {"Authorization": f"Bearer {ai21.api_key}", "Content-Type": "application/json"}
    payload = {
        "context": transcript,
        "question": question
    }
    endpoint = "https://api.ai21.com/studio/v1/experimental/answer"
    try:
        result = requests.post(endpoint, json = payload, headers=headers)
        result = result.json()
        answer = result["answer"]
        return answer
    except Exception as e:
        return {"status" : 0, "error": f"{type(e)} {e}"}

def createDoc(transcript):
    headers = {"Authorization": f"Bearer {ai21.api_key}", "Content-Type": "application/json"}
    payload = {
        "context": transcript,
        "question": "What are the tasks and people responsible for them?"
    }
    endpoint = "https://api.ai21.com/studio/v1/experimental/answer"
    try:
        summary = ai21.Summarize.execute(
            source=transcript,
            sourceType="TEXT"
        )
    except Exception as e:
        print(e)
        return {"status" : 0, "error": f"{type(e)} {e}"}
    
    summary = summary['summary']
    document = Document()
    document.add_heading('Meeting Minutes', 0)

    document.add_heading('Summary', level=1)
    document.add_paragraph(summary)
    document.add_paragraph("")
    answer = ask(transcript, "Give me a summary of the meeting in bullet points")
    document.add_paragraph(answer)

    document.add_heading('Responsibilities', level=1)
    answer = ask(transcript, "What are the tasks and people responsible for them?")
    document.add_paragraph(
        answer
    )

    document.save('minutes.docx')
    return document

@app.route('/scribe', methods=["POST"])
def scribe():
    if request.method == "POST":
        path = "../minutes.docx"
        url = request.get_json()['url']
        transcript = get_transcription(url)
        document = createDoc(transcript)
        #return ("OK", 200)
        #return send_file(path, as_attachment=True)
        return send_from_directory('../', 'minutes.docx', as_attachment=True)
    else:
        pass
    
@app.route('/clear', methods=["POST"])
def clear():
    if request.method == "POST":
        url = request.get_json()['url']
        r.delete(url)
        return ("OK", 200)
    else:
        pass

@app.route('/summarize', methods=["POST"])
def summarise():
    if request.method == "POST":
        #data = request.data
        #data = data.decode("utf-8")
        # data = eval(data) #data should only have one argument
        url = request.get_json()['url']
        transcript = get_transcription(url)
        try:
            summary = ai21.Summarize.execute(
                source=transcript,
                sourceType="TEXT"
            )
        except Exception as e:
            return (str(e), 400)

        return ({'summary':summary["summary"]}, 200)

    else:
        pass


@app.route('/append-transcript', methods=["POST"])
def appendTranscript():
    if request.method == 'POST':
        body = request.get_json()
        if 'url' in body and 'stream' in body:
            append_transcription(body['url'], body['stream'])
            return ('OK', 200)
        else:
            return ('Bad Request', 400)

@app.route('/get-transcript', methods=["POST"])
def getTranscript():
    if request.method == 'POST':
        body = request.get_json()
        if 'url' in body:
            value = get_transcription(body['url'])
            return (value, 200)
        else:
            return ('Bad Request', 400)


if __name__ == '__main__':
    # conn_dict = psycopg.conninfo.conninfo_to_dict(connection)
    app.run(debug=True)
